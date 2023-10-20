################################################################################ 
#####################      SCHEDULE DOCUMENTS LIBRARY     ###################### 
################################################################################
'''
A class, initialized by [date], which allows you to create rolls and bales Buggy 
schedules in the form of word documents and email them out to specified email 
addresses
'''
import os
import smtplib
import datetime
import calendar
from docx import Document
from email import encoders
from dotenv import load_dotenv
from tkinter import messagebox
from email.mime.base import MIMEBase
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import constants

class ScheduleDocuments:
    def __init__(self, date):
        load_dotenv()
        self.date = date

    def __createEqualGroups(self, people, n):
        '''
        Split people into n groups equally
        ''' 
        groups = []
        for i in range(n):
            start = int(i * len(people) / n)
            end = int((i + 1) * len(people) / n)
            groups.append(people[start:end])
        return list(reversed(sorted(groups, key=len)))

    def __validateDate(self):
        '''
        Validates that self.date is in the following format DD-MM-YYYY and
        and complies with other date standards such as 1 < MM < 12
        '''
        currentDateList = (str(datetime.date.today())).split("-")
        buggyDateList = self.date.split("-")
        
        try:
            for i in range(3):
                if not buggyDateList[i].isnumeric():
                    return False 
                if i == 2:
                    if len(buggyDateList[i]) != 4:
                        return False
                else:
                    if len(buggyDateList[i]) != 2:
                        return False
            if int(buggyDateList[0]) < 1 or  int(buggyDateList[0]) > 12:
                return False
            if int(buggyDateList[1]) < 1 or int(buggyDateList[1]) > 31:
                return False  
            if currentDateList[0] > buggyDateList[2]:
                return False
            elif currentDateList[0] < buggyDateList[2]:
                return True
            elif currentDateList[1] > buggyDateList[0]:
                return False
            elif currentDateList[1] < buggyDateList[0]:
                return True
            elif currentDateList[2] > buggyDateList[1]:
                return False
            else:
                return True
        except:
            return False

    def __getWeekday(self, offset=0):
        '''
        Gets the current weekday, optionally going forward (+) or backward (-)
        days based on offset
        ''' 
        try: 
            month, day, year = map(int, self.date.split('-'))
            dateObj = datetime.datetime(year, month, day)
            weekdayInt = dateObj.weekday() + offset
            res = calendar.day_name[weekdayInt]
        except Exception as e:
            messagebox.showerror("Error", "Invalid date: " + str(e))
        return res


    def __getPrevDate(self, offset=1):
        '''
        Gets the date [offset] days ago which defaults to 1. Return in form
        MM-DD-YYYY
        ''' 
        try:
            givenDate = datetime.datetime.strptime(self.date, "%m-%d-%Y")
            prevDate = givenDate - datetime.timedelta(days=offset)
            prevDateStr = prevDate.strftime("%m-%d-%Y")
        except Exception as e:
            messagebox.showerror("Error", "Invalid date: " + str(e))      
        return prevDateStr

    def __addRoleInformation(self, document, people, role):
        '''
        Adds a [role] section header and enumerated list of [people] to 
        [document]
        ''' 
        title = document.add_paragraph().add_run(role + ":")
        title.underline = True
        title.bold = True
        
        document.add_paragraph().add_run("Where and when to meet: TDB")
        individuals_string = ""
        for i, person in enumerate(people):
            if i == 0:
                individuals_string += f"\t{i+1}. {person}"
            else:
                individuals_string += f"\n\t{i+1}. {person}"
        
        document.add_paragraph().add_run(individuals_string)
    
    def __addPusherInformation(self, document, groups):
        '''
        Adds a pusher section header and formated display of [groups] to 
        [document]
        ''' 
        title = document.add_paragraph().add_run("Pushers:")
        title.underline = True
        title.bold = True
        
        document.add_paragraph().add_run("Where and when to meet: TBD")
        pushersString = ""
        for i in range(len(groups)):
            line = ", ".join(groups[i])
            if i == 0:
                pushersString += f"\t{i+1} | {line}"
            else:
                pushersString += f"\n\t{i+1} | {line}"
        
        document.add_paragraph().add_run(pushersString)

    def __createBalesDocument(self, bales, date, weekday):
        '''
        Create a docx for bales
        ''' 
        document = Document()
            
        document.add_heading("Bales-" + date + " (" + weekday + ")", 0)
        t1 = document.add_paragraph().add_run("Bales:")
        t1.underline = True
        t1.bold = True
        balesString = ""
        for i in range(len(bales)):
            if i == 0:
                balesString += bales[i]
            else:
                balesString = balesString + ", " + bales[i]
        document.add_paragraph().add_run(balesString)

        return document

    def __createRollsDocument(self, choicesDict, weekday):
        '''
        Create a docx for rolls
        ''' 
        document = Document()

        document.add_heading("Rolls-" + self.date + " (" + weekday + ")", 0)
        
        if choicesDict["barricaders"] != ["N/A"]:
            self.__addRoleInformation(
                document, choicesDict["barricaders"], "Barricaders"
            )
        
        if choicesDict["flaggers"] != ["N/A"]:
            self.__addRoleInformation(
                document, choicesDict["flaggers"], "Flaggers"
            )

        if choicesDict["pushers"] != ["N/A"]:
            groups = self.__createEqualGroups(
                choicesDict["pushers"], constants.PUSHERGROUPS
            )
            self.__addPusherInformation(document, groups)
        
        if choicesDict["videotimers"] != ["N/A"]:
            self.__addRoleInformation(
                document, choicesDict["videotimers"], "Video & Timers"
            )
        
        if choicesDict["backups"] != ["N/A"]:
            self.__addRoleInformation(
                document, choicesDict["backups"], "Backups"
            )

        document.add_page_break()

        t8 = document.add_paragraph().add_run("General Schedule:")
        t8.bold = True
        document.add_paragraph().add_run(constants.GENERALSCHEDULE)

        return document

    # create 
    def create(self, choicesDict):
        '''
        Use helpers to:
            1. check if date is valid
            2. get some useful date values such as week day and bales date
            3. create bales doc (if applicable) and save it to past folder
            4. create rolls doc and sace it to past folder
        We return None if error, [rolls] if no bales doc required and 
        [bales, rolls] if we require both bales and rolls docs
        ''' 
        if not self.__validateDate():
            messagebox.showerror(
                "Error", 
                "Please enter a valid future date in the form MM-DD-YYYY"
            )
            return None
        
        weekday = self.__getWeekday()
        prevWeekday = self.__getWeekday(offset=-1)
        prevDate = self.__getPrevDate()

        files = []
        # bales document
        if choicesDict["bales"] != ["N/A"]:
            balesDocument = self.__createBalesDocument(
                choicesDict["bales"], prevDate, prevWeekday
            )
            balesDocument.save("past/bales-" + prevDate + ".docx")
            files.append("past/bales-" + prevDate + ".docx")

        # Rolls document
        rollsDocument = self.__createRollsDocument(choicesDict, weekday)
        rollsDocument.save("past/rolls-" + self.date + ".docx")
        files.append("past/rolls-" + self.date + ".docx")

        return files

    def email(self, address, files):
        '''
        Sends an email to [address] with [files] attached from email account
        configured with .env file
        ''' 
        if not files:
            return None

        msg = MIMEMultipart()
        msg["From"] = os.getenv("EMAIL")
        msg["To"] = address
        msg["Subject"] = "Buggy-" + self.date
        msg.attach(MIMEText(""))
        for filename in files:
            attachment = open(filename, "rb")
            filename1 = filename.removeprefix("past/")
            part = MIMEBase("application", "octet-stream")
            part.set_payload(attachment.read())
            encoders.encode_base64(part)
            part.add_header("Content-Disposition",
            f"attachment; filename= {filename1}")
            msg.attach(part)
        msg = msg.as_string()
        try:
            server = smtplib.SMTP("smtp.outlook.com:587")
            server.ehlo()
            server.starttls()
            server.login(os.getenv("EMAIL"), os.getenv("PASSWORD"))
            server.sendmail(os.getenv("EMAIL"), address, msg)
            server.quit()
            messagebox.showinfo("Success", "The email sent successfully")
        except Exception as e:
            messagebox.showerror("Error", "Could not send email: " + str(e))
